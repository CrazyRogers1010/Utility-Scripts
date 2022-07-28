# Roster Generator

import os
import docx
import time
import openpyxl
import datetime
import shutil
import datetime
from pathlib import Path


os.chdir('/Users/christopherrogers/OneDrive - Collierville Schools/Cloud Downloads')

list_of_scoresheets = []
for file_name in os.listdir('/Users/christopherrogers/OneDrive - Collierville Schools/Cloud Downloads'):
    if file_name[0:16] == 'scoresheetReport':
        list_of_scoresheets.append(file_name)
most_recent_time = 0
most_recent_file = ''
now = datetime.datetime.now()

time_of_this_morning = time.time() - 3600*now.hour - 60*now.minute - 60*now.second

for scoresheet in list_of_scoresheets:
    print(Path(scoresheet))

    if os.stat(scoresheet)[8] > most_recent_time and os.stat(scoresheet)[8] > time_of_this_morning:
          most_recent_time = os.stat(scoresheet)[8]
          most_recent_file = scoresheet

if most_recent_file == '':
    print('No files from today. Download the full scoresheet from PowerSchool.')
else:
    print('The most recent scoresheet file is ' + most_recent_file)
    file = os.path.abspath(most_recent_file)
    print(file)
    wb = openpyxl.load_workbook(file)

    document = docx.Document()


    for sheet in wb.worksheets:
       
        list_of_students = []   
        max_row = sheet.max_row
        max_col = sheet.max_column

        for i in range(2,max_row):
            
            list_of_students.append(sheet.cell(row = i, column = 1).value)
        
        paragraph = document.add_paragraph(sheet.title + ' | '+ 'Mr. Rogers' + ' | ' + 'Rm. A531')

        table = document.add_table( rows=len(list_of_students)+1, cols=6)
        table.style = 'Table Grid'
        
        col_titles = ['Name', 'Present', 'Absent', 'Missing', 'Injured at FA', 'Injured, transported']

        for i in range(len(col_titles)):
            cell = table.cell(0, i)
            cell.text = col_titles[i]
            cell.paragraphs[0].runs[0].font.size = Pt(10)

        for i in range(len(list_of_students)):
            cell = table.cell(i+1, 0)
            cell.paragraphs[0].add_run(text=list_of_students[i])
            paragraph = cell.paragraphs[0]
            run = paragraph.runs
            font = run[0].font
            font.size = Pt(10)
            
        document.add_page_break()




    date = str(datetime.date.today())
    hour = str(datetime.datetime.now().hour)
    minute = str(datetime.datetime.now().minute)
    sec = str(datetime.datetime.now().second)


    os.chdir('/Users/christopherrogers/OneDrive - Collierville Schools/20-21/Classes/Rosters')
    shutil.copy(file, 'Rosters ' + str(datetime.date.today()) + ' ' + hour + '-' + minute + '-' + sec + '.xlsx')
    document.save('Rosters ' + str(datetime.date.today()) + ' ' + hour + '-' + minute + '-' + sec + '.docx')

    os.chdir('/Users/christopherrogers/Desktop')
    document.save('Rosters ' + str(datetime.date.today()) + ' ' + hour + '-' + minute + '-' + sec + '.docx')


    print('Done')

    
